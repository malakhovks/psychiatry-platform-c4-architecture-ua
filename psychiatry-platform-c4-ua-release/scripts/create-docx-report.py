#!/usr/bin/env python3
"""Create a comprehensive Ukrainian DOCX report for the C4 architecture package."""
from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Architecture_psychiatry_platform_C4_UA.docx"
REPORT_IMG = ROOT / "qa" / "report-images"
REPORT_IMG.mkdir(parents=True, exist_ok=True)

FONT = "Noto Sans"
FONT_MONO = "DejaVu Sans Mono"
NAVY = "17365D"
BLUE = "2F75B5"
LIGHT_BLUE = "D9EAF7"
PALE = "F2F6FA"
DARK = RGBColor(31, 50, 71)
MUTED = RGBColor(89, 100, 112)
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=85, bottom=70, end=85) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run_font(run, name=FONT, size: float | None = None, bold=None, color=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_para(paragraph, before=0, after=0, line=1.0, keep_next=False, keep_lines=False) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_next
    fmt.keep_together = keep_lines
    fmt.widow_control = True


def add_text(paragraph, text: str, *, size=9.0, bold=False, color=DARK, mono=False) -> None:
    run = paragraph.add_run(text)
    set_run_font(run, FONT_MONO if mono else FONT, size, bold, color)


def add_body(doc: Document, text: str, *, size=9.3, after=4, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    set_para(p, after=after, line=1.08)
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, size=size, bold=True)
        add_text(p, text[len(bold_prefix):], size=size)
    else:
        add_text(p, text, size=size)
    return p


def add_bullets(doc: Document, items: Iterable[str], *, size=9.1, compact=False):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_para(p, after=0 if compact else 2, line=1.05)
        add_text(p, item, size=size)


def add_heading(doc: Document, text: str, level=1, page_break_before=False):
    p = doc.add_paragraph(style=f"Heading {level}")
    if page_break_before:
        p.paragraph_format.page_break_before = True
    add_text(p, text, size={1: 17, 2: 13, 3: 10.5}.get(level, 10), bold=True,
             color=RGBColor(23, 54, 93) if level <= 2 else DARK)
    set_para(p, before=2 if level == 1 else 4, after=5 if level == 1 else 3, keep_next=True)
    return p


def add_page_break(doc: Document):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Сторінка ")
    set_run_font(run, FONT, 8, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_doc_defaults(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.05)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)
    section.header_distance = Cm(0.45)
    section.footer_distance = Cm(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(9.3)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in (("Heading 1", 17, NAVY), ("Heading 2", 13, BLUE), ("Heading 3", 10.5, NAVY)):
        st = styles[name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.space_before = Pt(4)
        st.paragraph_format.space_after = Pt(4)

    for style_name in ("List Bullet", "List Number"):
        st = styles[style_name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(9.1)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "Архітектура платформи психіатричних сценаріїв та алгоритмів"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in hp.runs:
        set_run_font(r, FONT, 8, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def optimize_image(src: Path, max_px=2600) -> Path:
    out = REPORT_IMG / src.name
    if out.exists() and out.stat().st_mtime >= src.stat().st_mtime:
        return out
    with Image.open(src) as im:
        im = im.convert("RGB")
        im.thumbnail((max_px, max_px), Image.Resampling.LANCZOS)
        # Preserve line art quality while significantly reducing package size.
        im.save(out, "PNG", optimize=True, compress_level=9)
    return out


def add_image_fit(doc: Document, path: Path, max_width_in: float, max_height_in: float, alt: str = ""):
    with Image.open(path) as im:
        w, h = im.size
    scale = min(max_width_in / w, max_height_in / h)
    width = Inches(w * scale)
    height = Inches(h * scale)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=3, keep_lines=True)
    run = p.add_run()
    inline = run.add_picture(str(path), width=width, height=height)
    try:
        doc_pr = inline._inline.docPr
        doc_pr.set("descr", alt)
        doc_pr.set("title", alt[:250])
    except Exception:
        pass
    return p


def add_kv_table(doc: Document, rows: list[tuple[str, str]], label_width=1.55, font_size=7.7):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        row = table.rows[i]
        keep_row_together(row)
        row.cells[0].width = Inches(label_width)
        row.cells[1].width = Inches(9.15)
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        if i % 2:
            set_cell_shading(row.cells[1], PALE)
        for c in row.cells:
            set_cell_margins(c, top=45, bottom=45, start=70, end=70)
        p0 = row.cells[0].paragraphs[0]
        set_para(p0, after=0, line=1.0, keep_lines=True)
        add_text(p0, label, size=font_size, bold=True, color=RGBColor.from_string(NAVY))
        p1 = row.cells[1].paragraphs[0]
        set_para(p1, after=0, line=1.0, keep_lines=True)
        add_text(p1, value, size=font_size, color=DARK)
    return table


def add_summary_table(doc: Document):
    data = [
        ("C1", "5", "Ландшафт і системні контексти"),
        ("C2", "6", "Повне та фокусні контейнерні подання"),
        ("C3", "13", "Компоненти ключових контейнерів"),
        ("L4", "2", "Контракти алгоритму та сесії"),
        ("DYN", "13", "Універсальні робочі процеси"),
        ("Deployment", "2", "Виробниче та edge/offline розгортання"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Рівень", "Кількість", "Покриття"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, text, size=8.8, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for idx, row_data in enumerate(data):
        cells = table.add_row().cells
        if idx % 2:
            for c in cells:
                set_cell_shading(c, PALE)
        for i, text in enumerate(row_data):
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
            add_text(p, text, size=8.5, bold=(i == 0), color=DARK)
    return table


def add_workflow_table(doc: Document):
    steps = [
        ("1", "Ініціація", "Мета звернення, роль, канал, контекст і допустимий профіль алгоритму."),
        ("2", "Ідентичність і доступ", "Автентифікація, мета доступу, роль, лікувальний контекст."),
        ("3", "Інформування і згода", "Версія повідомлення, обсяг згоди, окрема згода на медіа та ШІ."),
        ("4", "Відбір", "Критерії включення/виключення та негайні кризові ознаки."),
        ("5", "Збір даних", "Скарги, потреби, структуровані відповіді, медіа або біосигнали."),
        ("6", "Детермінований крок", "Типізація, валідація, розгалуження, наступний допустимий стан."),
        ("7", "Обчислення", "Бали, підшкали, пороги, правила неповних даних, червоні прапорці."),
        ("8", "Мультимодальна оцінка", "Контроль якості, аналіз модальностей, версії аналізаторів."),
        ("9", "Клінічна агрегація", "Пояснюваний проєкт результату, протоколу й траєкторії."),
        ("10", "Людський контроль", "Підтвердження, зміна або відхилення клінічно значущого проєкту."),
        ("11", "Маршрутизація", "Візит, телемедицина, кризова служба, контрольний строк і резервний канал."),
        ("12", "Супровід", "Нагадування, повторне опитування, вимірювання, пауза й відновлення."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [0.55, 2.15, 8.0]
    for i, text in enumerate(("№", "Етап", "Зміст")):
        cell = table.rows[0].cells[i]
        cell.width = Inches(widths[i])
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=50, bottom=50)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, text, size=8.2, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for idx, vals in enumerate(steps):
        row = table.add_row()
        keep_row_together(row)
        for i, text in enumerate(vals):
            c = row.cells[i]
            c.width = Inches(widths[i])
            if idx % 2:
                set_cell_shading(c, PALE)
            set_cell_margins(c, top=36, bottom=36, start=55, end=55)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
            add_text(p, text, size=7.5, bold=(i == 1), color=DARK)
    return table


def add_diagram_index(doc: Document, catalog: list[dict]):
    cols = 3
    rows = (len(catalog) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, item in enumerate(catalog):
        r = i % rows
        c = i // rows
        cell = table.cell(r, c)
        set_cell_margins(cell, top=45, bottom=45, start=70, end=70)
        if r % 2:
            set_cell_shading(cell, PALE)
        p = cell.paragraphs[0]
        set_para(p, after=0, line=1.0)
        add_text(p, f"{item['number']:02d}. ", size=7.7, bold=True, color=RGBColor.from_string(BLUE))
        add_text(p, item["title"], size=7.7, color=DARK)
    return table


def build_document() -> Path:
    catalog = json.loads((ROOT / "diagram-catalog.json").read_text(encoding="utf-8"))
    doc = Document()
    set_doc_defaults(doc)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(2.0)
    add_text(p, "АРХІТЕКТУРНИЙ ПАКЕТ", size=12, bold=True, color=RGBColor.from_string(BLUE))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, before=4, after=9, keep_lines=True)
    add_text(p, "Архітектура програмної платформи\nдля реалізації сценаріїв і алгоритмів\nу сфері психіатрії", size=25, bold=True, color=RGBColor.from_string(NAVY))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=14)
    add_text(p, "Повний набір C4-діаграм у Structurizr DSL та комплексний текстовий опис", size=12.5, color=MUTED)

    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    metrics = [("41", "діаграма"), ("424", "ідентифікаторів"), ("283", "статичних зв'язків"), ("13", "динамічних процесів")]
    for col, (value, label) in enumerate(metrics):
        set_cell_shading(table.cell(0, col), NAVY)
        set_cell_shading(table.cell(1, col), LIGHT_BLUE)
        set_cell_margins(table.cell(0, col), top=80, bottom=30)
        set_cell_margins(table.cell(1, col), top=30, bottom=80)
        p0 = table.cell(0, col).paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p0, value, size=18, bold=True, color=WHITE)
        p1 = table.cell(1, col).paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p1, label, size=8.5, bold=True, color=RGBColor.from_string(NAVY))
    merged = table.cell(2, 0).merge(table.cell(2, 3))
    set_cell_shading(merged, PALE)
    set_cell_margins(merged, top=70, bottom=70)
    p = merged.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "Мова діаграм: українська  •  Формат: Structurizr DSL  •  Дата оновлення: 21 червня 2026 року", size=8.8, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(0.8)
    add_text(p, "Дослідний архітектурний проєкт. Клінічне використання конкретного алгоритму можливе лише після формального зіставлення з авторитетним джерелом, клінічної валідації та погодження.", size=9.2, bold=True, color=DARK)

    # Executive summary
    add_heading(doc, "1. Керівне резюме", level=1, page_break_before=True)
    add_body(doc, "Платформа призначена для формалізації, версіонування, виконання та клінічного контролю психіатричних сценаріїв. Вона підтримує два взаємопов'язані контури: ШІ-асистент пацієнта та ШІ-асистент медичного працівника. Конкретні опитувальники й алгоритми підключаються як версійовані визначення, які виконує єдиний рушій.", size=10.0, after=6)
    add_summary_table(doc)
    add_heading(doc, "Ключове розділення відповідальності", level=2)
    add_bullets(doc, [
        "Детермінований контур виконує валідацію, переходи, формули, підшкали, пороги, червоні прапорці, кризові правила та маршрутизацію.",
        "Контур підтримки клінічних рішень агрегує результати, підбирає перевірені протоколи та формує пояснюваний проєкт траєкторії.",
        "Контур генеративного ШІ виконує транскрипцію, структуроване вилучення, RAG, адаптацію форми комунікації й формує лише проєкти тексту.",
        "Остаточне клінічно значуще рішення підтверджує, коригує або відхиляє уповноважений фахівець.",
    ], size=9.1)
    add_heading(doc, "Базові принципи", level=2)
    add_bullets(doc, [
        "Скринінг не є діагнозом; результат шкали подається з обмеженнями інтерпретації.",
        "Кризова логіка не залежить від LLM і має основний та резервний маршрут реагування.",
        "Активна сесія закріплена за конкретною незмінною версією алгоритму.",
        "Кожний результат має відтворюване походження: вхідні дані, правило, версія, час і суб'єкт дії.",
        "Низька якість або відсутність даних відображаються явно; система не заповнює прогалини вигаданими значеннями.",
    ], size=9.1, compact=True)

    # Source schema
    add_heading(doc, "2. Вихідна високорівнева схема та її відображення", level=1, page_break_before=True)
    src_img = optimize_image(ROOT / "sources" / "high-level-schema.jpg", max_px=1600)
    add_image_fit(doc, src_img, 7.2, 5.0, "Надана високорівнева схема контурів пацієнта і медичного працівника")
    add_kv_table(doc, [
        ("Контур пацієнта", "Скарги й потреби; опитувальники; супровід вимірювань; нагадування і календар; допомога із застосунком; підготовка до візиту; текст, голос, відео, аватар, VR і голос-в-текст; дозволена адаптація стилю, тембру та швидкості."),
        ("Контур фахівця", "Структурування скарг, опитувань, вимірювань і завдань; підбір протоколів; червоні прапорці; аналіз біосигналів, голосу, відео та зображень; Міннесотське кодування ЕКГ; проєкт траєкторії пацієнта."),
        ("Архітектурне рішення", "Обидва контури використовують спільний детермінований рушій, реєстр алгоритмів, клінічне сховище, аудит і контроль згод. Комунікаційний ШІ не виконує критичну клінічну логіку."),
        ("Трасування", "Повна матриця вимог і діаграм наведена у `docs/00-source-traceability.md`; конкретний клінічний зміст імпортується з контрольованого репозиторію джерел."),
    ], label_width=1.7, font_size=8.0)

    # Universal workflow
    add_heading(doc, "3. Універсальний робочий процес", level=1, page_break_before=True)
    add_body(doc, "Однакова структура вихідних алгоритмів реалізується метаданими, а не окремим програмним кодом. Нижче наведено інваріантний процес, який адаптується конфігурацією конкретного профілю.", size=9.6, after=5)
    add_workflow_table(doc)
    add_body(doc, "Кожний крок породжує аудиторську подію. Кризова ознака перериває звичайний workflow і переводить сесію до пріоритетного маршруту з підтвердженням отримання та резервною ескалацією.", size=9.2, after=0)

    # Contract overview
    add_heading(doc, "4. Формальна модель алгоритму", level=1, page_break_before=True)
    add_body(doc, "Універсальний контракт `AlgorithmDefinition` є агрегатом, який перетворює клінічний документ на виконувану, тестовану та відтворювану специфікацію. Сесія `ScenarioSession` фіксує стан конкретного проходження без прихованої зміни версії.", size=9.6, after=5)
    contract_rows = [
        ("Ідентичність", "ID, назва, власник, джерело, локаль, версія, статус, контрольна сума, дата чинності."),
        ("Застосовність", "Мета, популяція, критерії включення/виключення, протипоказання та дозволені ролі."),
        ("Згода", "Версія інформування, обов'язкові та опціональні дозволи, строк і правила відкликання."),
        ("Кроки", "Групи, питання, типи відповідей, валідація, обов'язковість, пропуск, повтор і завершення."),
        ("Логіка", "Граф станів, умови переходів, формули, підшкали, пороги і правила неповних даних."),
        ("Ризик", "Червоні прапорці, рівень критичності, максимальний строк, основний і резервний маршрут."),
        ("Вимірювання", "Модальність, пристрій, якість, оброблення, версія аналізатора та допустимість результату."),
        ("Результати", "Структурований результат, пояснення, межі інтерпретації, маршрутизація і follow-up."),
        ("Комунікація", "Канали, доступність, мова, стиль, тембр, швидкість та обмеження адаптації."),
        ("Контроль", "Набір тестів, рецензенти, цифрові погодження, журнал змін, відкликання і аудит."),
    ]
    add_kv_table(doc, contract_rows, label_width=1.6, font_size=8.0)
    add_heading(doc, "Підтримувані профілі", level=2)
    add_body(doc, "Архітектурний каталог містить точки підключення для PHQ-9, GAD-7, PQ-16, MDQ, CFQ, MoCA, AUDIT-C, DAST, оцінювання психічного статусу, кризового реагування, ЕКГ, голосу й відео. Повні тексти інструментів не дублюються у пакеті та мають проходити окрему перевірку прав використання.", size=9.2)

    # Index
    add_heading(doc, "5. Каталог діаграм", level=1, page_break_before=True)
    add_body(doc, "Наступні 41 сторінка містять окреме подання, його графічний перегляд і комплексний текстовий опис. Високороздільні SVG та PNG доступні в каталозі `previews/`.", size=9.5, after=6)
    add_diagram_index(doc, catalog)

    type_labels = {
        "systemLandscape": "C1 / ландшафт систем",
        "systemContext": "C1 / системний контекст",
        "container": "C2 / контейнери",
        "component": "C3 / компоненти",
        "custom": "L4 / логічний контракт",
        "dynamic": "Динамічне подання",
        "deployment": "Подання розгортання",
    }

    # One diagram per page.
    for item in catalog:
        add_heading(doc, f"{item['number'] + 5}. {item['title']}", level=1, page_break_before=True)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para(p, after=3)
        add_text(p, f"Ключ: {item['key']}  •  Тип: {type_labels.get(item['type'], item['type'])}  •  Елементів: {item['elementCount']}  •  Зв'язків/кроків: {item['relationshipOrStepCount']}", size=7.8, bold=True, color=MUTED)
        img = optimize_image(ROOT / item["png"], max_px=2800)
        add_image_fit(doc, img, 10.45, 3.35, item["title"])
        add_kv_table(doc, [
            ("Стислий зміст", item["description"]),
            ("Призначення", item["purpose"]),
            ("Як читати", item["reading"]),
            ("Ключові рішення", item["decisions"]),
            ("Безпека", item["safety"]),
            ("Аудиторія", item["audience"]),
        ], label_width=1.45, font_size=7.35)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_para(p, before=2, after=0)
        add_text(p, f"Повнороздільне подання: {item['svg']}", size=6.9, color=MUTED, mono=True)

    # Closing sections
    add_heading(doc, "47. Безпека, приватність і клінічне управління", level=1, page_break_before=True)
    add_bullets(doc, [
        "Zero-trust для зовнішніх інтеграцій: взаємна автентифікація, короткоживучі облікові дані, перевірка контракту й дозволеної мети.",
        "Розділення операційного стану, клінічних записів, алгоритмів, медіа, бази знань, аналітики та незмінного аудиту.",
        "Згода є версійованим ресурсом; окремі дозволи застосовуються до медіа, біосигналів, зовнішніх моделей і вторинного використання даних.",
        "Правила доступу поєднують роль, лікувальний контекст, мету, чутливість, згоду і стан кризового випадку.",
        "Генеративна модель викликається тільки через контрольований шлюз із мінімізацією даних, RAG за погодженими джерелами та післягенераційною перевіркою.",
        "Кризові події мають підтвердження доставки, максимальний строк реакції, повторну спробу, резервний канал і видимий статус для фахівця.",
        "Будь-який машинний аналіз біосигналу або медіа зберігає якість, версію моделі, ознаки, обмеження й факт клінічного перегляду.",
    ], size=9.2)
    add_heading(doc, "Мінімальний аудиторський запис", level=2)
    add_kv_table(doc, [
        ("Хто/коли", "Ідентичність суб'єкта або сервісу, роль, час, кореляційний ідентифікатор і пристрій."),
        ("Що", "Подія, ресурс, версія алгоритму, вхідні дані або контрольна сума, результат і статус."),
        ("Чому", "Мета доступу, правило, джерело протоколу, причина зміни або клінічне обґрунтування."),
        ("Як", "Сервіс, версія коду або моделі, політика, канал, результат контролю якості й запобіжники."),
    ], label_width=1.35, font_size=8.2)

    add_heading(doc, "48. Рекомендована послідовність реалізації", level=1, page_break_before=True)
    roadmap = [
        ("Етап 1", "Архітектурне ядро", "Ідентичність, згода, API, реєстр алгоритмів, базовий рушій, аудит, PHQ-9/GAD-7 як пілотні профілі."),
        ("Етап 2", "Клінічний контур", "Кабінет фахівця, червоні прапорці, підтримка рішень, траєкторія, FHIR-інтеграція."),
        ("Етап 3", "Мультимодальність", "Пристрої, quality gate, ЕКГ, голос, відео, медичні зображення, контроль моделей."),
        ("Етап 4", "Масштабування", "Шина подій, HA, edge/offline, аналітика, багатозакладність і політики зберігання."),
        ("Етап 5", "Клінічна експлуатація", "Валідація, навчання персоналу, SOP, моніторинг безпеки, постмаркетинговий контроль і регулярне відкликання версій."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(("Етап", "Фокус", "Основний результат")):
        set_cell_shading(table.rows[0].cells[i], NAVY)
        set_cell_margins(table.rows[0].cells[i])
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, h, size=8.5, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for idx, vals in enumerate(roadmap):
        row = table.add_row()
        keep_row_together(row)
        for i, val in enumerate(vals):
            if idx % 2:
                set_cell_shading(row.cells[i], PALE)
            set_cell_margins(row.cells[i])
            p = row.cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
            add_text(p, val, size=8.3, bold=(i == 1), color=DARK)
    add_heading(doc, "Критерії готовності профілю алгоритму", level=2)
    add_bullets(doc, [
        "задокументоване авторитетне джерело та права використання;",
        "формальна схема пройшла валідацію;",
        "контрольні приклади покривають звичайні, граничні, неповні та кризові випадки;",
        "клінічний власник і незалежний рецензент погодили версію;",
        "маршрути ескалації перевірено в тестовому середовищі;",
        "моніторинг, аудит, відкликання й резервне відновлення налаштовані.",
    ], size=9.0, compact=True)

    add_heading(doc, "49. Статус валідації та склад файлів", level=1, page_break_before=True)
    add_body(doc, "Під час формування пакета виконано локальну статичну перевірку узгодженості: 424 ідентифікатори, 283 статичні відношення та 41 унікальний ключ подання; перевірено JSON Schema, каталог JSON, усі YAML-приклади та наявність 41 SVG і 41 PNG. Для презентаційних C1.01, C1.03 і C1.04 метрично перевірено 75 текстових блоків: усі написи розміщено всередині відповідних фігур або текстових областей. Офіційний Structurizr parser не запускався в цьому середовищі через відсутність Docker/локального бінарного пакета; для цього додано `scripts/validate.sh`.", size=9.6, after=6)
    add_kv_table(doc, [
        ("Основний DSL", "`workspace.dsl` — модульний workspace з `!include`; `workspace-single-file.dsl` — переносна однофайлова версія."),
        ("Модель", "`model/` — ролі, системи, контейнери, компоненти, контракти, deployment-вузли й відношення."),
        ("Подання", "`views/` — C1, C2, C3, L4, 13 DYN, 2 deployment та єдині стилі."),
        ("Документація", "`docs/` — огляд, принципи, workflow, каталог усіх діаграм, контракти, безпека, розгортання й roadmap."),
        ("Рішення", "`adrs/` — 10 ADR щодо детермінованого ядра, версій, людського контролю, кризового контуру, згод, сховищ, ШІ, подій, якості та edge."),
        ("Контракти", "`schemas/algorithm-definition.schema.json`; `examples/algorithm-template.yaml`, `algorithm-catalog.yaml`, `test-case-template.yaml`."),
        ("Перегляди", "`previews/index.html`, 41 SVG, 41 PNG і 41 DOT-файл."),
        ("Автоматизація", "`scripts/build-single-file.py`, `lint-workspace.py`, `render-previews.py`, `render_c1_presentation.py`, `create-c1-presentation.js`, `validate.sh`, `export.sh`."),
    ], label_width=1.5, font_size=8.0)
    add_heading(doc, "Висновок", level=2)
    add_body(doc, "Запропонована архітектура створює керовану основу для широкого набору психіатричних сценаріїв без дублювання програмного коду. Її головна властивість — відокремлення клінічно критичної детермінованої логіки від допоміжного генеративного ШІ, із повною версійністю, простежуваністю, людським контролем та fail-safe кризовим контуром.", size=10.0, after=0)

    # Metadata
    props = doc.core_properties
    props.title = "Архітектура програмної платформи психіатричних сценаріїв та алгоритмів"
    props.subject = "C4 model, Structurizr DSL, psychiatry, clinical algorithms"
    props.author = "Архітектурний проєкт TNMU"
    props.keywords = "C4, Structurizr, психіатрія, алгоритми, ШІ, клінічна безпека"
    props.comments = "Комплексний українськомовний опис 41 C4-подання."

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build_document()
    print(path)
